"""
Build presentacion_vargas.html from encuesta vargas.docx (static, self-contained).
Matches SIGNOS-style layout from presentacion_encuesta; adds scroll-in animations.
"""
from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.table import Table

DOCX_PATH = Path(r"c:\Users\mz\Downloads\encuesta vargas.docx")
OUT_HTML = Path(__file__).resolve().parent / "presentacion_vargas.html"
PULSE_HREF = "presentacion_atacama_pulse_ia.html"
ASSETS_PREFIX = "assets/"

DECK_NAV_HTML = """  <nav class="deck-nav-fixed" aria-label="Navegación entre presentaciones">
    <a class="deck-nav-link deck-nav-prev" href="presentacion_encuesta.html">← Anterior</a>
    <a class="deck-nav-link deck-nav-hub" href="index.html#deck-catalogo">Índice</a>
    <a class="deck-nav-link deck-nav-next" href="{pulse}">Siguiente →</a>
  </nav>
"""

PCT_RE = re.compile(r"([\d.,]+)\s*%?\s*$")

# Entities and quoted labels to emphasize in qualitative blocks (longest-first match).
_QUAL_KEY_TERMS: tuple[str, ...] = (
    "Servicio Local de Educación Pública Atacama",
    "Servicio Local de Educación Pública",
    "Fondo Nacional de Desarrollo Regional",
    "Subcomisaría Pedro León Gallo",
    "Centro Oncológico Atacama",
    "Corte de Apelaciones de Copiapó",
    "Corporación Ciahn",
    "Fundación Comprometidos",
    "Atacama Cultural",
    "Congreso Futuro 2026",
    "Desierto Creativo",
    "King Long",
    "Miguel Vargas Correa",
    "Miguel Vargas",
    "Partido Socialista",
    "Plan Atacama",
    "Red Copiapó",
    "SernamEG",
    "Electromovilidad",
    "Inteligencia Artificial",
    "Chañaral",
    "Copiapó",
    "Caldera",
    "Atacama",
    "SEREMI",
    "Intendente",
    "SLEP",
    "GORE",
    "FNDR",
)

_CONTINUATION_START = re.compile(
    r"^(podría|podrían|este documento)\b|^(la conclusión|el escenario|la percepción general)\b",
    re.I,
)

_TRANSITION_PARA_START = re.compile(
    r"^(Además|Sin embargo|La conclusión|El informe|No obstante|Aun así|"
    r"Sin embargo también|Es decir|Pese al|Pero también|Sin embargo,? también)\b",
    re.I,
)


def esc(s: str) -> str:
    return html.escape(s, quote=True)


def merge_span_intervals(spans: list[tuple[int, int]]) -> list[tuple[int, int]]:
    """Union overlapping intervals so overlapping highlights merge into one <mark>."""
    if not spans:
        return []
    spans_sorted = sorted(spans)
    out: list[tuple[int, int]] = [(spans_sorted[0][0], spans_sorted[0][1])]
    for s, e in spans_sorted[1:]:
        ls, le = out[-1]
        if s <= le:
            out[-1] = (ls, max(le, e))
        else:
            out.append((s, e))
    return out


def highlight_qualitative(text: str) -> str:
    """Wrap key terms, percentages, and quotes in <mark class='qual-hl'>; escapes safely."""
    if not text.strip():
        return ""
    spans: list[tuple[int, int]] = []
    for term in sorted(_QUAL_KEY_TERMS, key=len, reverse=True):
        for m in re.finditer(re.escape(term), text, flags=re.IGNORECASE):
            spans.append(m.span())

    quote_rx = r'«[^»]{2,180}»|“[^”]{2,180}”|"[^"]{2,160}"|\[[^\]]{2,80}\]'
    for m in re.finditer(quote_rx, text):
        spans.append(m.span())

    for m in re.finditer(r"\b\d+[,.]\d+\s*%", text):
        spans.append(m.span())

    merged = merge_span_intervals(spans)
    pos = 0
    parts: list[str] = []
    for s, e in merged:
        if s < pos:
            continue
        parts.append(esc(text[pos:s]))
        parts.append(f'<mark class="qual-hl">{esc(text[s:e])}</mark>')
        pos = e
    parts.append(esc(text[pos:]))
    return "".join(parts)


def format_card_title(title: str) -> str:
    """Card heading with optional percentage chip at the end."""
    m = re.search(r"^(.*?)(\s*[—\-]\s*[\d.,]+\s*%)\s*$", title.strip())
    if m:
        head, chip = m.group(1), m.group(2).strip()
        return f"{highlight_qualitative(head)} <span class='qual-pct-chip'>{esc(chip)}</span>"
    return highlight_qualitative(title)


def split_sentences_es(s: str) -> list[str]:
    chunks = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÑ¿¡\"«“])", s.strip())
    return [c.strip() for c in chunks if c.strip()]


def format_narrative_body(body_paras: list[str]) -> str:
    """Use bullet lists when several paragraphs or several sentences in one block."""
    paras = [p.strip() for p in body_paras if p and p.strip()]
    if not paras:
        return ""

    if len(paras) >= 2:
        lis = "".join(
            f'<li class="qual-li anim-item">{highlight_qualitative(p)}</li>' for p in paras
        )
        return f'<ul class="qual-bullets narr-bullets">{lis}</ul>'

    only = paras[0]
    sents = split_sentences_es(only)
    if len(sents) >= 3:
        lis = "".join(
            f'<li class="qual-li anim-item">{highlight_qualitative(s)}</li>' for s in sents
        )
        return f'<ul class="qual-bullets narr-bullets">{lis}</ul>'

    return f'<p class="narr-prose">{highlight_qualitative(only)}</p>'


def format_radar_body_lines(lines: list[str]) -> str:
    """Turn short runs after a colon lead-in into bullet lists; keep long paragraphs as <p>."""
    raw = [x.strip() for x in lines if x.strip()]
    if not raw:
        return ""

    parts: list[str] = []
    expecting_bullets = False
    bullet_buf: list[str] = []
    short_max = 158

    def flush_bullets() -> None:
        nonlocal bullet_buf
        if len(bullet_buf) >= 2:
            lis = "".join(
                f'<li class="qual-li anim-item">{highlight_qualitative(b)}</li>' for b in bullet_buf
            )
            parts.append(f'<ul class="qual-bullets radar-bullets">{lis}</ul>')
        elif bullet_buf:
            parts.append(f'<p class="radar-p">{highlight_qualitative(bullet_buf[0])}</p>')
        bullet_buf = []

    i = 0
    while i < len(raw):
        s = raw[i]
        i += 1

        if _TRANSITION_PARA_START.match(s) and not s.rstrip().endswith(":"):
            flush_bullets()
            expecting_bullets = False
            parts.append(f'<p class="radar-p">{highlight_qualitative(s)}</p>')
            continue

        if s.rstrip().endswith(":") and len(s) < 220:
            flush_bullets()
            expecting_bullets = True
            parts.append(f'<p class="radar-lead">{highlight_qualitative(s)}</p>')
            continue

        if expecting_bullets:
            st = s.strip()
            if len(st) <= short_max and not _CONTINUATION_START.match(st):
                bullet_buf.append(s)
                continue
            flush_bullets()
            expecting_bullets = False

        parts.append(f'<p class="radar-p">{highlight_qualitative(s)}</p>')

    flush_bullets()
    return "\n".join(parts)


def parse_pct_cell(s: str) -> float | None:
    t = (s or "").strip().replace("%", "").replace(",", ".")
    try:
        return float(t)
    except ValueError:
        return None


def is_tail_label(label: str) -> bool:
    l = label.lower()
    return bool(
        re.search(r"no\s+sabe|no\s+responde|ns\s*/\s*nr", l)
        or re.match(r"^otros?$", l.strip(), re.I)
        or re.match(r"^total$", l.strip(), re.I)
    )


def sort_by_pct_rows(
    rows: list[tuple[str, str]], *, tail_last: bool = True
) -> list[tuple[str, str]]:
    def sort_key(item: tuple[str, str]) -> tuple[int, float, str]:
        label, pct_s = item
        p = parse_pct_cell(pct_s) or 0.0
        if tail_last and is_tail_label(label):
            return (1, p, label)
        return (0, -p, label)

    return sorted(rows, key=sort_key)


def table_pair_rows(table: Table, col0: int = 0, col1: int = 1) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) <= col1:
            continue
        out.append((cells[col0], cells[col1]))
    return out


def hbar_block(rows: list[tuple[str, str]], *, dual: bool = False) -> str:
    rows = sort_by_pct_rows([(a, b) for a, b in rows if not re.match(r"^total$", a, re.I)])
    parts: list[str] = []
    for label, pct_s in rows:
        if re.match(r"^total$", label, re.I):
            continue
        w = parse_pct_cell(pct_s) or 0.0
        parts.append(
            f'<div class="hbar-row anim-item"><span class="hbar-label">{esc(label)}</span>'
            f'<div class="hbar-track"><span class="hbar-fill" style="--w:{w:.2f};"></span></div>'
            f'<span class="hbar-pct">{esc(pct_s.strip())}</span></div>'
        )
    return '<div class="hbar-block">' + "".join(parts) + "</div>"


def comparison_table(table: Table) -> str:
    header = [c.text.strip() for c in table.rows[0].cells]
    body_rows: list[str] = []
    for row in table.rows[1:]:
        vals = [c.text.strip() for c in row.cells]
        if vals[0].lower().startswith("total"):
            body_rows.append("<tr><td><em>" + esc(vals[0]) + "</em></td>" + "".join(
                f'<td class="num">{esc(v)}</td>' for v in vals[1:]
            ) + "</tr>")
        else:
            body_rows.append(
                "<tr><td>"
                + esc(vals[0])
                + "</td>"
                + "".join(f'<td class="num">{esc(v)}</td>' for v in vals[1:])
                + "</tr>"
            )
    thead = "<tr>" + "".join(f"<th>{esc(h)}</th>" for h in header) + "</tr>"
    return f'<table class="data-table anim-item"><thead>{thead}</thead><tbody>{"".join(body_rows)}</tbody></table>'


def comparison_bars(table: Table, col2024: int = 1, col2026: int = 2) -> str:
    parts: list[str] = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        label = cells[0]
        if re.match(r"^total$", label, re.I):
            continue
        a = parse_pct_cell(cells[col2024]) or 0.0
        b = parse_pct_cell(cells[col2026]) or 0.0
        parts.append(
            f'<div class="cmp-row anim-item"><div class="cmp-label">{esc(label)}</div>'
            f'<div class="cmp-bars"><span class="cmp-y">2024</span>'
            f'<div class="cmp-track"><span class="hbar-fill cmp-a" style="--w:{a:.2f};"></span></div>'
            f'<span class="cmp-p">{esc(cells[col2024])}%</span>'
            f'<span class="cmp-y">2026</span>'
            f'<div class="cmp-track"><span class="hbar-fill cmp-b" style="--w:{b:.2f};"></span></div>'
            f'<span class="cmp-p">{esc(cells[col2026])}%</span></div></div>'
        )
    return '<div class="cmp-wrap">' + "".join(parts) + "</div>"


def prose_card(title: str, body_paras: list[str]) -> str:
    body_html = format_narrative_body(body_paras)
    return (
        f'<article class="narr-card anim-item"><h4 class="narr-card-title">{format_card_title(title)}</h4>'
        f'<div class="narr-card-body">{body_html}</div></article>'
    )


def extract_numbered_pairs(sl: list[tuple[str, str]]) -> list[tuple[str, list[str]]]:
    """Group '1. Title...' with following paragraphs until next numbered item."""
    blocks: list[tuple[str, list[str]]] = []
    i = 0
    while i < len(sl):
        line = sl[i][0]
        if re.match(r"^\d+\.\s+", line.strip()):
            title = line
            i += 1
            body_chunks: list[str] = []
            while i < len(sl):
                nxt = sl[i][0]
                if re.match(r"^\d+\.\s+", nxt.strip()):
                    break
                body_chunks.append(nxt)
                i += 1
            blocks.append((title, body_chunks))
        else:
            i += 1
    return blocks


def load_doc() -> tuple[Document, list[tuple[str, str]]]:
    doc = Document(str(DOCX_PATH))
    paras: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        t = " ".join((p.text or "").split())
        if t:
            paras.append((t, getattr(p.style, "name", "") or ""))
    return doc, paras


def build() -> None:
    doc, paras = load_doc()
    tables = doc.tables

    meth_kicker = paras[0][0] if paras else ""
    meth_sub = paras[1][0] if len(paras) > 1 else ""

    # Principales problemas
    probs = sort_by_pct_rows(table_pair_rows(tables[0]))

    gov_reg_title = paras[6][0]  # approximate by scan
    for t, _ in paras:
        if "Gobierno Regional de Atacama" in t:
            gov_reg_title = t
            break
    gov_gob_title = "Evaluación de la Gestión del Gobernador Regional Miguel Vargas Correa"
    for t, _ in paras:
        if "Gobernador Regional Miguel Vargas" in t:
            gov_gob_title = t
            break

    attr_heading = paras[11][0] if len(paras) > 11 else "Atributos"
    attr_intro = paras[12][0] if len(paras) > 12 else ""
    buses_title = paras[16][0] if len(paras) > 16 else "Buses eléctricos"

    pos_start = None
    neg_start = None
    for idx, (t, _) in enumerate(paras):
        if "Distribución de Ejes Tematicos" in t or "Ejes Tematicos" in t:
            pos_start = idx + 1
        if "Distribución de Atributos Negativos" in t:
            neg_start = idx + 1
            break
    pos_slice = paras[pos_start:neg_start - 1] if pos_start and neg_start else []
    neg_slice_end = None
    for idx, (t, _) in enumerate(paras):
        if "Radar Social" in t or "Analisi cualitativo" in t:
            neg_slice_end = idx
            break
    neg_slice = paras[neg_start:neg_slice_end] if neg_start and neg_slice_end else []

    pos_blocks = extract_numbered_pairs(pos_slice)
    neg_blocks = extract_numbered_pairs(neg_slice)

    radar_idx = None
    for idx, (t, _) in enumerate(paras):
        if "Radar Social" in t or "Analisi cualitativo" in t:
            radar_idx = idx
            break
    radar_paras = paras[radar_idx + 1 :] if radar_idx is not None else []

    # Build narrative HTML for radar: group by lines starting with digit+dot at line start
    radar_sections: list[tuple[str, list[str]]] = []
    cur_title = ""
    cur_lines: list[str] = []
    for t, _ in radar_paras:
        if re.match(r"^\d+\.\s+", t.strip()):
            if cur_title or cur_lines:
                radar_sections.append((cur_title, cur_lines))
            cur_title = t
            cur_lines = []
        else:
            cur_lines.append(t)
    if cur_title or cur_lines:
        radar_sections.append((cur_title, cur_lines))

    nav_ids = [
        ("portada", "Portada"),
        ("metodologia", "Metodología"),
        ("problemas", "Problemas"),
        ("gore", "GORE"),
        ("gobernador", "Gobernador"),
        ("atributos", "Atributos"),
        ("judicial", "Judicial"),
        ("buses", "Buses"),
        ("positivos", "Ejes +"),
        ("negativos", "Ejes −"),
        ("radar", "Radar"),
        ("pulse-ia", "Pulse IA"),
    ]
    nav_html = "".join(f'<a href="#{i}">{l}</a>' for i, l in nav_ids)

    slides: list[str] = []

    slides.append(
        f"""
<header class="cover-hero slide-reveal" id="portada" aria-label="Portada">
  <div class="cover-bg" style="background-image:url('{ASSETS_PREFIX}cover-signos.png')"></div>
</header>
"""
    )

    slides.append(
        f"""
<section class="slide slide-metodologia slide-reveal" id="metodologia">
  <span class="pill">Metodología</span>
  <h2>Encuesta regional — Antecedentes</h2>
  <p class="lead anim-item">{esc(meth_kicker)} · {esc(meth_sub)}</p>
</section>
"""
    )

    # Table problems
    probs_rows_html = "".join(
        f"<tr><td>{esc(a)}</td><td class='num'>{esc(b)}</td></tr>" for a, b in probs
    )
    slides.append(
        f"""
<section class="slide slide-reveal" id="problemas">
  <span class="pill">Resultados</span>
  <h2>Principales problemas de la región</h2>
  <div class="slide-body">
    <div class="panel">{hbar_block(probs)}</div>
    <div class="panel scroll-x"><table class="data-table"><thead><tr><th>Tema</th><th>%</th></tr></thead><tbody>{probs_rows_html}</tbody></table></div>
  </div>
</section>
"""
    )

    slides.append(
        f"""
<section class="slide slide-reveal" id="gore">
  <span class="pill">Evaluación</span>
  <h2>{esc(gov_reg_title)}</h2>
  <div class="slide-body one-col">
    {comparison_table(tables[1])}
    {comparison_bars(tables[1])}
  </div>
</section>
"""
    )

    slides.append(
        f"""
<section class="slide slide-reveal" id="gobernador">
  <span class="pill">Evaluación</span>
  <h2>{esc(gov_gob_title)}</h2>
  <div class="slide-body one-col">
    {comparison_table(tables[2])}
    {comparison_bars(tables[2])}
  </div>
</section>
"""
    )

    dim_rows = "".join(
        f"<tr><td>{esc(r[0])}</td><td class='num'>{esc(r[1])}</td></tr>"
        for r in table_pair_rows(tables[3])
        if not re.match(r"dimensión|total", r[0], re.I)
    )
    slides.append(
        f"""
<section class="slide slide-reveal" id="atributos">
  <span class="pill">Escala 1–7</span>
  <h2>{esc(attr_heading)}</h2>
  <p class="lead anim-item">{esc(attr_intro[:520])}{"…" if len(attr_intro) > 520 else ""}</p>
  <div class="panel scroll-x"><table class="data-table"><thead><tr><th>Dimensión</th><th>Nota</th></tr></thead><tbody>{dim_rows}</tbody></table></div>
</section>
"""
    )

    cred_rows = table_pair_rows(tables[4])
    cred_for_bar = [(a, b) for a, b in cred_rows if not re.match(r"total|nivel", a, re.I)]
    slides.append(
        f"""
<section class="slide slide-reveal" id="judicial">
  <span class="pill">Percepción</span>
  <h2>Credibilidad ante la situación judicial</h2>
  <div class="slide-body">
    <div class="panel">{hbar_block(cred_for_bar)}</div>
    <div class="panel scroll-x"><table class="data-table"><thead><tr><th>Nivel</th><th>%</th></tr></thead><tbody>{"".join(f"<tr><td>{esc(a)}</td><td class='num'>{esc(b)}</td></tr>" for a,b in cred_rows)}</tbody></table></div>
  </div>
</section>
"""
    )

    bus_rows = table_pair_rows(tables[5])
    bus_bar = [(a, b) for a, b in bus_rows if not re.match(r"categoría|total", a, re.I)]
    slides.append(
        f"""
<section class="slide slide-reveal" id="buses">
  <span class="pill">Infraestructura</span>
  <h2>{esc(buses_title)}</h2>
  <div class="slide-body">
    <div class="panel">{hbar_block(bus_bar)}</div>
    <div class="panel scroll-x"><table class="data-table"><thead><tr><th>Categoría</th><th>%</th></tr></thead><tbody>{"".join(f"<tr><td>{esc(a)}</td><td class='num'>{esc(b)}</td></tr>" for a,b in bus_rows)}</tbody></table></div>
  </div>
</section>
"""
    )

    pos_cards = "".join(prose_card(a, b) for a, b in pos_blocks if a)
    slides.append(
        f"""
<section class="slide slide-reveal" id="positivos">
  <span class="pill">Cualitativo</span>
  <h2>Distribución de ejes temáticos — atributos positivos</h2>
  <div class="narr-grid">{pos_cards}</div>
</section>
"""
    )

    neg_cards = "".join(prose_card(a, b) for a, b in neg_blocks if a)
    slides.append(
        f"""
<section class="slide slide-reveal" id="negativos">
  <span class="pill">Cualitativo</span>
  <h2>Distribución de atributos negativos</h2>
  <div class="narr-grid">{neg_cards}</div>
</section>
"""
    )

    radar_html_parts: list[str] = []
    for title, lines in radar_sections:
        if not title and not lines:
            continue
        body = format_radar_body_lines(lines)
        radar_html_parts.append(
            f'<div class="radar-block anim-item"><h3>{highlight_qualitative(title)}</h3><div class="radar-body">{body}</div></div>'
        )
    slides.append(
        f"""
<section class="slide slide-reveal" id="radar">
  <span class="pill">Radar Social</span>
  <h2>Análisis cualitativo</h2>
  <div class="radar-wrap">{"".join(radar_html_parts)}</div>
</section>
"""
    )

    slides.append(
        f"""
<section class="slide slide-reveal slide-pulse-cta" id="pulse-ia">
  <span class="pill">Comunicación e IA</span>
  <h2>Programa comunicacional automatizado · Atacama Pulse AI</h2>
  <p class="lead anim-item">Complemento estratégico a este informe: plataforma de <strong>inteligencia social</strong> hiperlocal, interacción comunitaria asistida por IA y crecimiento orgánico multi-red (sin foco exclusivo en pauta).</p>
  <p class="anim-item" style="font-size:.8rem;color:#3d475c;margin:.45rem 0 0;">Incluye validación respecto de algoritmos, arquitectura por módulos (radar, grupos, interacción, viral, growth, mapa de influencia), proyección de alcance regional y gestión de riesgos.</p>
  <p class="anim-item pulse-cta-actions">
    <a class="pulse-btn" href="{esc(PULSE_HREF)}">Abrir presentación Atacama Pulse AI (4 láminas)</a>
  </p>
</section>
"""
    )

    body_inner = "\n".join(slides)
    deck_nav_block = DECK_NAV_HTML.replace("{pulse}", esc(PULSE_HREF))

    html_out = f"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Encuesta Vargas — Región de Atacama</title>
  <link rel="preconnect" href="https://fonts.googleapis.com" />
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin />
  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;600;700&family=Montserrat:wght@600;700&display=swap" rel="stylesheet" />
  <link rel="stylesheet" href="{ASSETS_PREFIX}deck-nav.css" />
  <style>
    :root {{
      --signos-navy: #004a99;
      --signos-purple: #8a56c3;
      --signos-cyan: #29e3f1;
      --signos-deep: #003671;
      --dk: #1e2d44;
      --surface: #fff;
      --bg: linear-gradient(165deg, #e8f3fb 0%, #eef1f8 50%, #e4e9f2 100%);
      --shadow: 0 6px 18px rgba(0, 54, 113, 0.12);
      --radius: 10px;
      --maxw: 920px;
    }}
    @media (prefers-reduced-motion: reduce) {{
      .slide-reveal, .anim-item {{ animation: none !important; opacity: 1 !important; transform: none !important; transition: none !important; }}
      .slide-reveal:not(.is-visible) .hbar-fill:not(.cmp-a):not(.cmp-b),
      .slide-reveal:not(.is-visible) .hbar-fill.cmp-a,
      .slide-reveal:not(.is-visible) .hbar-fill.cmp-b {{ width: min(100%, calc(var(--w, 100) * 1%)) !important; }}
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; font-family: "DM Sans", system-ui, sans-serif; color: var(--dk); background: var(--bg); font-size: clamp(13px, 2.2vw, 15px); line-height: 1.52; }}
    .sticky-head {{ position: sticky; top: 0; z-index: 80; background: rgba(255,255,255,.97); backdrop-filter: blur(10px); border-bottom: 1px solid rgba(0,74,153,.12); }}
    .brand-bar {{ text-align: center; padding: .35rem 1rem; border-bottom: 1px solid rgba(0,74,153,.08); }}
    .brand-bar img {{ height: clamp(26px,4vw,34px); width: auto; max-width: min(220px,42vw); object-fit: contain; }}
    .brand-bar a {{ display: inline-block; line-height: 0; }}
    nav.toc .inner {{ max-width: var(--maxw); margin: 0 auto; padding: .45rem 1rem; display: flex; flex-wrap: wrap; gap: .35rem; align-items: center; font-size: .8rem; }}
    nav.toc a {{ padding: .2rem .55rem; border-radius: 999px; background: rgba(138,86,195,.12); color: var(--dk); font-weight: 600; text-decoration: none; }}
    nav.toc a:hover {{ background: rgba(41,227,241,.22); }}
    main {{ max-width: var(--maxw); margin: 0 auto; padding: .85rem .75rem 2.5rem; }}
    .cover-hero {{ position: relative; width: 100%; max-width: 900px; margin: 0 auto; aspect-ratio: 16/9; max-height: min(48vh,360px); min-height: 200px; border-radius: 0 0 var(--radius) var(--radius); overflow: hidden; box-shadow: var(--shadow); scroll-margin-top: 6.5rem; }}
    .cover-bg {{ position: absolute; inset: 0; background-color: var(--signos-navy); background-position: center; background-size: contain; background-repeat: no-repeat; }}
    .slide {{ background: var(--surface); border-radius: var(--radius); box-shadow: var(--shadow); padding: .75rem .9rem; margin-bottom: 1rem; scroll-margin-top: 6.5rem; max-width: 100%; overflow-x: hidden; }}
    .slide-reveal {{ opacity: 0; transform: translateY(16px); transition: opacity .55s ease, transform .55s ease; }}
    .slide-reveal.is-visible {{ opacity: 1; transform: translateY(0); }}
    .slide-reveal.is-visible .anim-item {{ animation: fadeInUp .5s ease backwards; }}
    .slide-reveal.is-visible .anim-item:nth-child(1) {{ animation-delay: .05s; }}
    .slide-reveal.is-visible .anim-item:nth-child(2) {{ animation-delay: .1s; }}
    .slide-reveal.is-visible .anim-item:nth-child(3) {{ animation-delay: .15s; }}
    .slide-reveal.is-visible .anim-item:nth-child(4) {{ animation-delay: .2s; }}
    .slide-reveal.is-visible .anim-item:nth-child(5) {{ animation-delay: .25s; }}
    .slide-reveal.is-visible .anim-item:nth-child(n+6) {{ animation-delay: .3s; }}
    .slide-reveal.is-visible .qual-li:nth-child(1) {{ animation-delay: .03s; }}
    .slide-reveal.is-visible .qual-li:nth-child(2) {{ animation-delay: .07s; }}
    .slide-reveal.is-visible .qual-li:nth-child(3) {{ animation-delay: .11s; }}
    .slide-reveal.is-visible .qual-li:nth-child(4) {{ animation-delay: .15s; }}
    .slide-reveal.is-visible .qual-li:nth-child(5) {{ animation-delay: .19s; }}
    .slide-reveal.is-visible .qual-li:nth-child(n+6) {{ animation-delay: .23s; }}
    .slide-reveal.is-visible .hbar-row:nth-child(1) {{ animation-delay: .05s; }}
    .slide-reveal.is-visible .hbar-row:nth-child(2) {{ animation-delay: .1s; }}
    .slide-reveal.is-visible .hbar-row:nth-child(3) {{ animation-delay: .15s; }}
    .slide-reveal.is-visible .hbar-row:nth-child(4) {{ animation-delay: .2s; }}
    .slide-reveal.is-visible .hbar-row:nth-child(5) {{ animation-delay: .25s; }}
    .slide-reveal.is-visible .hbar-row:nth-child(n+6) {{ animation-delay: .3s; }}
    @keyframes fadeInUp {{
      from {{ opacity: 0; transform: translateY(10px); }}
      to {{ opacity: 1; transform: translateY(0); }}
    }}
    .pill {{ display: inline-block; font-size: .62rem; font-weight: 700; letter-spacing: .08em; text-transform: uppercase; color: #fff; background: linear-gradient(90deg, var(--signos-purple), var(--signos-cyan)); padding: .12rem .45rem; border-radius: 999px; }}
    h2 {{ margin: .25rem 0 0; font-size: clamp(.95rem, 2.4vw, 1.12rem); line-height: 1.3; }}
    h3 {{ font-size: .95rem; color: var(--signos-deep); margin: 0 0 .5rem; }}
    h4 {{ margin: 0 0 .35rem; font-size: .88rem; color: var(--signos-purple); }}
    .lead {{ color: #4b5569; font-size: .84rem; margin: .5rem 0 0; }}
    .slide-body {{ display: grid; gap: .65rem; margin-top: .65rem; }}
    @media (min-width: 800px) {{ .slide-body:not(.one-col) {{ grid-template-columns: 1fr 1fr; }} }}
    .one-col {{ grid-template-columns: 1fr; }}
    .panel {{ min-width: 0; }}
    .scroll-x {{ overflow-x: auto; }}
    .data-table {{ width: 100%; border-collapse: collapse; font-size: .78rem; table-layout: fixed; }}
    .data-table th, .data-table td {{ border: 1px solid rgba(0,74,153,.12); padding: .28rem .38rem; text-align: left; vertical-align: top; word-wrap: break-word; }}
    .data-table th {{ background: rgba(0,74,153,.1); font-weight: 700; }}
    .data-table td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
    .hbar-block {{ display: flex; flex-direction: column; gap: .35rem; }}
    .hbar-row {{ display: grid; grid-template-columns: minmax(0,1fr) minmax(0,1.5fr) auto; gap: .35rem; align-items: center; }}
    .hbar-label {{ font-size: .75rem; min-width: 0; }}
    .hbar-track {{ height: 8px; border-radius: 999px; background: rgba(0,54,113,.1); overflow: hidden; }}
    .hbar-fill {{ display: block; height: 100%; border-radius: 999px; background: linear-gradient(90deg, var(--signos-purple), var(--signos-cyan)); width: min(100%, calc(var(--w, 0) * 1%)); min-width: 2px; transition: width .75s cubic-bezier(.2,.8,.2,1); }}
    .slide-reveal:not(.is-visible) .hbar-fill:not(.cmp-a):not(.cmp-b) {{ width: 2px !important; }}
    .slide-reveal:not(.is-visible) .hbar-fill.cmp-a, .slide-reveal:not(.is-visible) .hbar-fill.cmp-b {{ width: 2px !important; }}
    .hbar-pct {{ font-size: .72rem; font-weight: 600; font-variant-numeric: tabular-nums; }}
    .cmp-wrap {{ display: flex; flex-direction: column; gap: .6rem; margin-top: .5rem; }}
    .cmp-row {{ border: 1px solid rgba(0,74,153,.12); border-radius: 8px; padding: .45rem .5rem; background: rgba(255,255,255,.9); }}
    .cmp-label {{ font-weight: 600; font-size: .82rem; margin-bottom: .35rem; }}
    .cmp-bars {{ display: grid; grid-template-columns: auto 1fr auto; gap: .25rem .4rem; align-items: center; font-size: .72rem; }}
    .cmp-y {{ font-weight: 600; color: var(--signos-deep); opacity: .85; }}
    .cmp-track {{ height: 7px; border-radius: 999px; background: rgba(0,54,113,.08); overflow: hidden; }}
    .cmp-a {{ background: rgba(138,86,195,.85) !important; }}
    .cmp-b {{ background: rgba(41,227,241,.95) !important; }}
    .cmp-p {{ font-variant-numeric: tabular-nums; text-align: right; }}
    .narr-grid {{ display: grid; gap: .65rem; margin-top: .65rem; }}
    @media (min-width: 750px) {{ .narr-grid {{ grid-template-columns: repeat(2, minmax(0,1fr)); }} }}
    .narr-card {{ border: 1px solid rgba(138,86,195,.22); border-radius: var(--radius); padding: .65rem .75rem; background: linear-gradient(135deg, rgba(138,86,195,.06), rgba(41,227,241,.06)); }}
    .narr-card-title {{ line-height: 1.35; }}
    .narr-card-title mark.qual-hl {{ background: rgba(138,86,195,.28); }}
    mark.qual-hl {{
      background: linear-gradient(120deg, rgba(138,86,195,.34), rgba(41,227,241,.31));
      color: inherit;
      font-weight: 600;
      padding: .08em .16em;
      border-radius: 4px;
    }}
    .qual-pct-chip {{
      display: inline-block;
      margin-left: .3rem;
      font-size: .72em;
      font-weight: 700;
      color: var(--signos-deep);
      background: rgba(255,255,255,.92);
      border: 1px solid rgba(138,86,195,.42);
      padding: .12em .45em;
      border-radius: 999px;
      vertical-align: middle;
      letter-spacing: .02em;
    }}
    .narr-card-body {{ font-size: .8rem; color: #3d475c; }}
    .narr-card-body .narr-prose {{ margin: 0; }}
    .qual-bullets {{ margin: 0; padding-left: 1.12rem; list-style-position: outside; }}
    .qual-bullets li {{
      margin: .4rem 0;
      padding-left: .08rem;
    }}
    .qual-bullets li::marker {{ color: var(--signos-purple); font-weight: 700; }}
    .narr-bullets li {{ font-size: .8rem; line-height: 1.46; }}
    .radar-wrap {{ margin-top: .65rem; display: flex; flex-direction: column; gap: .9rem; }}
    .radar-block {{ padding: .65rem; border-radius: var(--radius); border-left: 4px solid var(--signos-purple); background: rgba(0,74,153,.04); }}
    .radar-body .radar-p {{ margin: 0 0 .55rem; font-size: .8rem; }}
    .radar-body .radar-lead {{ margin: 0 0 .4rem !important; font-size: .8rem; font-weight: 600; color: var(--signos-deep); }}
    .radar-body .radar-bullets {{ margin: 0 0 .55rem; }}
    .radar-body mark.qual-hl {{ font-weight: 600; }}
    .radar-block h3 mark.qual-hl {{ background: rgba(0,74,153,.14); }}
    .slide-metodologia {{ border-top: 3px solid var(--signos-cyan); }}
    .slide-pulse-cta {{
      border-top: 3px solid var(--signos-purple);
      background: linear-gradient(145deg, rgba(138,86,195,.08), rgba(41,227,241,.06));
    }}
    .pulse-cta-actions {{ margin: .75rem 0 0; }}
    .pulse-btn {{
      display: inline-flex;
      align-items: center;
      padding: .5rem .95rem;
      border-radius: 999px;
      font-weight: 700;
      font-size: .82rem;
      text-decoration: none;
      color: #fff !important;
      background: linear-gradient(90deg, var(--signos-purple), var(--signos-cyan));
      box-shadow: 0 4px 14px rgba(138,86,195,.32);
      transition: transform .2s ease, box-shadow .2s ease;
    }}
    .pulse-btn:hover {{ transform: translateY(-1px); box-shadow: 0 6px 18px rgba(0,74,153,.2); }}
    .pulse-btn:focus-visible {{ outline: 2px solid var(--signos-deep); outline-offset: 3px; }}
    footer {{ text-align: center; padding: 1.25rem .75rem; color: #6b7284; font-size: .78rem; }}
  </style>
</head>
<body class="has-deck-nav">
  <div class="sticky-head">
    <div class="brand-bar">
      <a href="#portada"><img src="{ASSETS_PREFIX}logo-signos.png" alt="SIGNOS analytics" decoding="async" /></a>
    </div>
    <nav class="toc" aria-label="Índice"><div class="inner"><strong style="margin-right:.35rem;color:var(--dk)">Ir a:</strong>{nav_html}</div></nav>
  </div>
  <main>
    {body_inner}
  </main>
  <footer>Documento fuente: encuesta vargas.docx · Presentación web estática. Complemento: <a href="{esc(PULSE_HREF)}">Atacama Pulse AI</a>.</footer>
  <script>
(function() {{
  const reduce = window.matchMedia('(prefers-reduced-motion: reduce)').matches;
  if (reduce) {{
    document.querySelectorAll('.slide-reveal').forEach(function(el) {{ el.classList.add('is-visible'); }});
    return;
  }}
  const io = new IntersectionObserver(function(entries) {{
    entries.forEach(function(e) {{
      if (e.isIntersecting) {{
        e.target.classList.add('is-visible');
        io.unobserve(e.target);
      }}
    }});
  }}, {{ threshold: 0.12, rootMargin: '0px 0px -5% 0px' }});
  document.querySelectorAll('.slide-reveal').forEach(function(el) {{ io.observe(el); }});
}})();
  </script>
{deck_nav_block}
</body>
</html>
"""

    OUT_HTML.write_text(html_out, encoding="utf-8")
    print(f"Wrote {OUT_HTML}")


if __name__ == "__main__":
    build()
